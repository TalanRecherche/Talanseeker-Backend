"""Created on Wed Aug 30 14:28:28 2023

@author: agarc

"""
import docx

from app.core.cv_information_retrieval.abc_reader import ABCReader


class DOCXReader(ABCReader):
    @staticmethod
    def read_text(file_path: str) -> str | None:
        """Read all text from a word (docs) file.

        Args:
        ----
        file_path (str): The path of the word file.

        Returns:
        -------
        str: The extracted text from the word file.

        Dependencies:
        import docx
        """
        # load document
        doc = docx.Document(file_path)
        # instantiate return
        full_text = []
        # parse through paragraphs
        for paragraph in doc.paragraphs:
            paragraph.text.strip()
            # if paragraph_text:
            full_text.append(paragraph.text)

        # parse through tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_text = paragraph.text.strip()
                        # if paragraph_text:
                        full_text.append(paragraph_text)

        # add linebreaks to each elements of the line
        full_text = [line + "\n" for line in full_text]

        # join each line
        string_text = "".join(full_text)

        # remove non ascci character
        string_text = string_text.encode("latin-1", errors="ignore").decode("latin-1")

        # remove trailing spaces and linebreaks
        string_text = string_text.rstrip("\n")
        string_text = string_text.strip()

        # return only if not empty
        if string_text:
            return string_text
        else:
            return None
